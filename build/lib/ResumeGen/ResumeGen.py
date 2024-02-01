from docx import Document
# from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_user_input():
    try:
        # Collect information from the user
        name = input("Enter your full name: ")
        phone = input("Enter your phone number: ")
        email = input("Enter your email address: ")

        # Basic error handling
        if not name or not phone or not email:
            print("Error: Please provide at least your name, phone number, and email address.")
            exit()

        # Other sections
        address = input("Enter your address: ")
        summary = input("Enter a summary or objective statement: ")
        education = input("Enter your educational background: ")
        skills = input("Enter your skills, separated by commas: ")
        experience = input("Enter your work experience: ")
        achievements = input("Enter your achievements or projects: ")
        hobbies = input("Enter your hobbies: ")
        languages = input("Enter languages you speak: ")
        references = input("Enter references: ")
        linkedin = input("Enter your LinkedIn profile URL: ")
        portfolio = input("Enter a link to your online portfolio: ")
    except Exception as e:
        print(f"\033[31m Error:{e}\033[0m")
        pass

    return {
        "Name": name,
        "Address": address,
        "Phone": phone,
        "Email": email,
        "Summary": summary,
        "Education": education,
        "Skills": skills,
        "Experience": experience,
        "Achievements": achievements,
        "Hobbies": hobbies,
        "Languages": languages,
        "References": references,
        "LinkedIn": linkedin,
        "Portfolio": portfolio,
    }


def create_docx(data, filename="Resume.docx"):
    # Create a new Word document
    doc = Document()

    # Title
    doc.add_heading('Resume',
                    0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    try:
        # Add personal information
        doc.add_heading('Personal Information', level=1)
        for key, value in data.items():
            if key not in ["Education", "Skills", "Experience", "Achievements", "References", "Hobbies", "Summary"] and value:
                p = doc.add_paragraph(style='BodyText')
                p.add_run(key + ": ").bold = True
                p.add_run(value)

        # Add summary
        if data["Summary"]:
            doc.add_heading('Summary Objective', level=1)
            doc.add_paragraph(data["Summary"])

        # Add education
        if data["Education"]:
            doc.add_heading('Education', level=1)
            doc.add_paragraph(data["Education"])

        # Add skills
        if data["Skills"]:
            doc.add_heading('Skills', level=1)
            doc.add_paragraph(data["Skills"])

        # Add work experience
        if data["Experience"]:
            doc.add_heading('Work Experience', level=1)
            doc.add_paragraph(data["Experience"])

        # Add achievements or projects
        if data["Achievements"]:
            doc.add_heading('Achievements/Projects', level=1)
            doc.add_paragraph(data["Achievements"])

        # Add hobbies
        if data["Hobbies"]:
            doc.add_heading('Hobbies', level=1)
            doc.add_paragraph(data["Hobbies"])

        # Add languages
        # if data["Languages"]:
            # doc.add_heading('Languages', level=1)
            # doc.add_paragraph(data["Languages"])

        # Add references
        if data["References"]:
            doc.add_heading('References', level=1)
            doc.add_paragraph(data["References"])
            print("\033[35mFinalizing...\033[0m")

        # Save the document as DOCX
        doc.save(filename)
        print(f"\033[32m Saving Resume as {filename}...>\033[34mDONE\033[0m")
    except Exception as e:
        print(f"\033[31mError:{e}\033[0m")


def main():
    try:
        user_data = get_user_input()
        create_docx(user_data)
    except Exception:
        pass


if __name__ == "__main__":
    # Create DOCX
    main()
